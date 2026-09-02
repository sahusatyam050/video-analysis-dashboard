import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx_report(output_dir, task_id=None, complaint_id=None):
    """
    Generates an editable Word document (evidence_report.docx) containing the complete forensic report and proof images.
    """
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, "evidence_report.docx")

    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Document Header
    title = doc.add_heading("FORENSIC VIDEO EVIDENCE REPORT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(26, 26, 46)

    # Subtitle
    sub = doc.add_paragraph(f"Task ID: #{task_id or 'N/A'}  |  Complaint ID: {complaint_id or 'N/A'}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.runs[0]
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Executive Summary Section
    h1 = doc.add_heading("1. Executive Summary & Verdict", level=1)
    h1.runs[0].font.color.rgb = RGBColor(26, 26, 46)

    txt_path = os.path.join(output_dir, "final_summary.txt")
    if os.path.exists(txt_path):
        with open(txt_path, "r", encoding="utf-8") as f:
            for line in f:
                line_str = line.strip()
                if line_str:
                    p = doc.add_paragraph(line_str)
                    p.paragraph_format.space_after = Pt(4)
    else:
        doc.add_paragraph("[Source summary text file not found]")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Segment Evidence & Proof Screenshots Section
    h2 = doc.add_heading("2. Segment Proof Screenshots & AI Analysis", level=1)
    h2.runs[0].font.color.rgb = RGBColor(26, 26, 46)

    seg_json_path = os.path.join(output_dir, "segment_verdicts.json")
    if os.path.exists(seg_json_path):
        with open(seg_json_path, "r", encoding="utf-8") as f:
            segments = json.load(f)

        for seg in segments:
            idx = seg.get("segment_index", 1)
            start = seg.get("start_time", 0.0)
            end = seg.get("end_time", 0.0)
            ai_sum = seg.get("ai_summary", "General navigation.")
            proof_path = seg.get("proof_frame")

            p_seg = doc.add_paragraph()
            r_head = p_seg.add_run(f"Segment #{idx} ({start:.2f}s – {end:.2f}s)")
            r_head.bold = True
            r_head.font.size = Pt(12)
            r_head.font.color.rgb = RGBColor(0, 102, 204)

            p_ai = doc.add_paragraph()
            r_lbl = p_ai.add_run("AI Summary: ")
            r_lbl.bold = True
            p_ai.add_run(ai_sum)

            if proof_path and os.path.exists(proof_path):
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(proof_path, width=Inches(3.5))
                except Exception as e:
                    doc.add_paragraph(f"[Image load error: {e}]")
            else:
                doc.add_paragraph("[Proof Frame Image Unavailable]")

            doc.add_paragraph().paragraph_format.space_after = Pt(12)
    else:
        doc.add_paragraph("[No segment verdicts found]")

    doc.save(out_path)
    return out_path

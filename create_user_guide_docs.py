import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def generate_docx():
    md_path = "PROJECT_USER_GUIDE.md"
    docx_path = "PROJECT_USER_GUIDE.docx"
    if not os.path.exists(md_path):
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Custom styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    for line in lines:
        line_s = line.strip()
        if not line_s or line_s.startswith("```") or line_s == "---":
            continue
            
        if line_s.startswith("# "):
            p = doc.add_heading(level=1)
            run = p.add_run(line_s[2:])
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            run.font.size = Pt(18)
            run.bold = True
        elif line_s.startswith("## "):
            p = doc.add_heading(level=2)
            run = p.add_run(line_s[3:])
            run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
            run.font.size = Pt(14)
            run.bold = True
        elif line_s.startswith("### "):
            p = doc.add_heading(level=3)
            run = p.add_run(line_s[4:])
            run.font.color.rgb = RGBColor(0x03, 0x69, 0xA1)
            run.font.size = Pt(12)
            run.bold = True
        elif line_s.startswith("- ") or line_s.startswith("* "):
            clean_t = line_s[2:].replace("**", "")
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(clean_t)
        else:
            clean_t = line_s.replace("**", "")
            p = doc.add_paragraph()
            run = p.add_run(clean_t)
            
    doc.save(docx_path)
    print(f"✅ Generated {docx_path}")

def generate_pdf():
    md_path = "PROJECT_USER_GUIDE.md"
    pdf_path = "PROJECT_USER_GUIDE.pdf"
    if not os.path.exists(md_path):
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    pdf_doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle', parent=styles['Heading1'], fontSize=18, leading=22, textColor=colors.HexColor('#0F172A'), spaceAfter=10
    )
    h2_style = ParagraphStyle(
        'DocH2', parent=styles['Heading2'], fontSize=13, leading=16, textColor=colors.HexColor('#0284C7'), spaceBefore=12, spaceAfter=6
    )
    h3_style = ParagraphStyle(
        'DocH3', parent=styles['Heading3'], fontSize=11, leading=14, textColor=colors.HexColor('#0369A1'), spaceBefore=8, spaceAfter=4
    )
    body_style = ParagraphStyle(
        'DocBody', parent=styles['Normal'], fontSize=9.5, leading=13.5, textColor=colors.HexColor('#334155'), spaceAfter=5
    )
    bullet_style = ParagraphStyle(
        'DocBullet', parent=styles['Normal'], fontSize=9.5, leading=13.5, textColor=colors.HexColor('#334155'), leftIndent=15, spaceAfter=3
    )
    
    story = []
    
    in_code_block = False
    for line in lines:
        line_s = line.strip()
        if line_s.startswith("```"):
            in_code_block = not in_code_block
            continue
        if not line_s or line_s == "---" or in_code_block:
            continue
            
        # Escape XML chars for ReportLab
        clean_text = line_s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # Simple markdown bold replace
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean_text)
        
        if line_s.startswith("# "):
            story.append(Paragraph(clean_text[2:], title_style))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E1"), spaceAfter=10))
        elif line_s.startswith("## "):
            story.append(Paragraph(clean_text[3:], h2_style))
        elif line_s.startswith("### "):
            story.append(Paragraph(clean_text[4:], h3_style))
        elif line_s.startswith("- ") or line_s.startswith("* "):
            story.append(Paragraph(f"• {clean_text[2:]}", bullet_style))
        else:
            story.append(Paragraph(clean_text, body_style))
            
    pdf_doc.build(story)
    print(f"✅ Generated {pdf_path}")

if __name__ == "__main__":
    generate_docx()
    generate_pdf()
